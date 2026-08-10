# -*- coding: utf-8 -*-
"""
VeFa - Tez Denetleme, Düzeltme ve Akademik Çeviri Modülü
Bu modül yüklenen .docx, .md veya .txt tez dosyalarını okur, denetler, düzeltir
ve isteğe bağlı olarak Türkçe, Arapça (العربية) veya İngilizce dillerine akademik olarak çevirir.
"""

import io
import logging
import os
import re
from typing import Dict, Any, Tuple, Optional, Callable
import docx

from engine.academic_engine import sanitize_turkish_text
from engine.turkish_prompts import LANG_CONFIGS

logger = logging.getLogger(__name__)


def _extract_footnotes_map(doc) -> Dict[str, str]:
    """
    Reads word/footnotes.xml (if present) and returns {footnote_id: footnote_text}.
    Skips the separator/continuation-separator entries (ids -1 and 0) that
    every Word document has by default and that carry no real text.
    """
    footnotes_map = {}
    try:
        footnotes_part = None
        for rel in doc.part.rels.values():
            if rel.reltype.endswith("/footnotes"):
                footnotes_part = rel.target_part
                break
        if footnotes_part is None:
            return footnotes_map

        from docx.oxml.ns import qn
        if hasattr(footnotes_part, "element"):
            root = footnotes_part.element
        else:
            from docx.oxml import parse_xml
            root = parse_xml(footnotes_part.blob)
        for fn in root.findall(qn('w:footnote')):
            fid = fn.get(qn('w:id'))
            if fid in ('-1', '0'):
                continue
            texts = [t.text or '' for t in fn.iter(qn('w:t'))]
            note_text = ''.join(texts).strip()
            if fid and note_text:
                footnotes_map[fid] = note_text
    except Exception as e:
        logger.warning(f"Dipnotlar okunamadı (footnotes.xml erişilemedi): {e}")
    return footnotes_map


def _paragraph_text_with_footnotes(paragraph, footnotes_map: Dict[str, str]) -> str:
    """
    Like paragraph.text, but inserts a [[fn:...]] marker at every point where
    the paragraph has a real Word footnote reference — so footnotes survive
    the extract -> audit/translate -> re-render round trip instead of
    silently vanishing (python-docx's paragraph.text ignores footnotes
    entirely, since they live in a separate part of the document).
    """
    from docx.oxml.ns import qn
    parts = []
    for run_el in paragraph._p.findall(qn('w:r')):
        for t in run_el.findall(qn('w:t')):
            parts.append(t.text or '')
        for ref in run_el.findall(qn('w:footnoteReference')):
            fid = ref.get(qn('w:id'))
            note_text = footnotes_map.get(fid)
            if note_text:
                parts.append(f"[[fn:{note_text}]]")
    return ''.join(parts)


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extract text content from uploaded .docx, .md, or .txt bytes.
    Real Word footnotes are preserved as [[fn:...]] markers so they survive
    the round trip through Mod 2 and can be re-rendered as real footnotes
    again by turkish_docx_generator.py.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".docx":
        doc = docx.Document(io.BytesIO(file_bytes))
        footnotes_map = _extract_footnotes_map(doc)
        paragraphs = []
        for p in doc.paragraphs:
            text_with_fn = _paragraph_text_with_footnotes(p, footnotes_map) if footnotes_map else p.text
            if text_with_fn.strip():
                paragraphs.append(text_with_fn.strip())
        for t in doc.tables:
            for row in t.rows:
                row_txt = [c.text.strip() for c in row.cells]
                if any(row_txt):
                    paragraphs.append("| " + " | ".join(row_txt) + " |")
        if footnotes_map:
            logger.info(f"{len(footnotes_map)} gerçek Word dipnotu [[fn:...]] işareti olarak korundu.")
        return "\n\n".join(paragraphs)

    elif ext in [".md", ".txt"]:
        try:
            return file_bytes.decode("utf-8-sig")
        except UnicodeDecodeError:
            return file_bytes.decode("cp1254", errors="ignore")

    else:
        raise ValueError(f"Desteklenmeyen dosya uzantısı: {ext}")


def translate_headings(text: str, target_lang: str) -> str:
    """
    Translate standard academic section headings to target language (tr, ar, en).
    """
    if target_lang not in LANG_CONFIGS:
        return text

    target_cfg = LANG_CONFIGS[target_lang]

    heading_map = {
        # Table of Contents
        r'^#?\s*(الفهرس والمحتويات|Alfurş ve İçerik|الفهرس|TABLE OF CONTENTS|İÇİNDEKİLER).*$': f"# {target_cfg['toc_title']}",

        # Turkish / English / Arabic Chapter 1
        r'^#?\s*(?:1\.\s*BÖLÜM|BÖLÜM 1|Bölüm 1|CHAPTER 1|الفصل الأول|GİRİŞ|INTRODUCTION).*$': f"# {target_cfg['intro']}",
        # Chapter 2
        r'^#?\s*(?:2\.\s*BÖLÜM|BÖLÜM 2|Bölüm 2|CHAPTER 2|الفصل الثاني|LİTERATÜR TARAMASI|LITERATURE REVIEW).*$': f"# {target_cfg['lit_review']}",
        # Chapter 3
        r'^#?\s*(?:3\.\s*BÖLÜM|BÖLÜM 3|Bölüm 3|CHAPTER 3|الفصل الثالث|YÖNTEM|METHODOLOGY).*$': f"# {target_cfg['methodology']}",
        # Chapter 4
        r'^#?\s*(?:4\.\s*BÖLÜM|BÖLÜM 4|Bölüm 4|CHAPTER 4|الفصل الرابع|BULGULAR|RESULTS).*$': f"# {target_cfg['results']}",
        # Chapter 5
        r'^#?\s*(?:5\.\s*BÖLÜM|BÖLÜM 5|Bölüm 5|CHAPTER 5|الفصل الخامس|TARTIŞMA|DISCUSSION).*$': f"# {target_cfg['discussion']}",
        # Chapter 6
        r'^#?\s*(?:6\.\s*BÖLÜM|BÖLÜM 6|Bölüm 6|CHAPTER 6|الفصل السادس|SONUÇ|CONCLUSION).*$': f"# {target_cfg['conclusion']}",
        # References / Abstract
        r'^#?\s*(KAYNAKÇA|BIBLIYOGRAFYA|REFERENCES|BIBLIOGRAPHY|المصادر والمراجع|المراجع).*$': f"# {target_cfg['references']}",
        r'^#?\s*(ÖZET|ABSTRACT|الملخص).*$': f"# {target_cfg['abstract']}",
    }

    for pattern, repl in heading_map.items():
        text = re.sub(pattern, repl, text, flags=re.MULTILINE | re.IGNORECASE)

    # Keywords replacement
    text = re.sub(r'\b(Anahtar Kelimeler|Keywords|الكلمات المفتاحية):\s*', f"{target_cfg['keywords']}: ", text, flags=re.IGNORECASE)

    return text


def audit_and_fix_thesis(
    text: str,
    llm_func: Optional[Any] = None,
    metadata: Optional[Dict[str, Any]] = None,
    target_language: str = "same",  # 'same', 'tr', 'ar', 'en'
    fix_language: bool = True,
    fix_yok_formatting: bool = True,
    fix_citations: bool = True,
    progress_callback: Optional[Callable[[float, str, str], None]] = None
) -> Tuple[str, Dict[str, Any]]:
    """
    Audit input text for academic & formatting issues, perform corrections and translation,
    and return cleaned markdown and detailed Audit Report.
    """
    audit_findings = []
    corrections_made = []
    cleaned_text = text

    # 1. Audit & strip Meta-notes / AI fluff (always applied, regardless of
    # fix_language, so leftover "Word count: 500" style notes never survive)
    meta_notes = re.findall(r'(Word count:\s*\d+|Target word count:\s*.*?|Here is the text.*?:)', cleaned_text, flags=re.IGNORECASE)
    if meta_notes:
        audit_findings.append(f"⚠️ **Yapay Zeka Meta Notları Bulundu ({len(meta_notes)} adet):** Metin sonunda kalan kelime sayısı notları tespit edildi.")
        cleaned_text = re.sub(r'^\s*(Word count|Target word count|Here is the text.*?):.*$', '', cleaned_text, flags=re.MULTILINE | re.IGNORECASE)
        corrections_made.append("✅ Metin sonundaki 'Word count' ve yapay zeka sunuş notları tamamen temizlendi.")

    # 2. Audit Heading Structure
    bold_headings = re.findall(r'^\s*\*\*(\d+(?:\.\d+)*\.?)\s+([^*#\n]+)\*\*', cleaned_text, flags=re.MULTILINE)
    if bold_headings:
        audit_findings.append(f"⚠️ **Biçimlendirilmemiş Kalın Başlıklar ({len(bold_headings)} adet):** Hiyerarşik `#` işareti içermeyen başlıklar bulundu.")
        if fix_yok_formatting:
            def _fix_h(match):
                num, title = match.group(1), match.group(2).strip()
                dots = num.rstrip('.').count('.')
                hashes = '#' * min(3, dots + 1)
                return f"{hashes} {num} {title}"
            cleaned_text = re.sub(r'^\s*\*\*(\d+(?:\.\d+)*\.?)\s+([^*#\n]+)\*\*', _fix_h, cleaned_text, flags=re.MULTILINE)
            corrections_made.append("✅ Tüm düz başlıklar akademik `# 1. BÖLÜM` ve `## 1.1.` Markdown başlık stillerine çevrildi.")
        else:
            corrections_made.append("ℹ️ YÖK başlık düzenine çevirme seçeneği kapalı olduğundan başlıklar orijinal halinde bırakıldı.")

    # 3. Clean text / Sanitize (Only if fix_language is True)
    if fix_language:
        cleaned_text = sanitize_turkish_text(cleaned_text, language=target_language if target_language != 'same' else 'tr')

    # 4. Citations Adjustment (Only if fix_citations is True)
    if fix_citations:
        et_als = re.findall(r'\bet al\.?(?=[\s,;:\)]|$)', cleaned_text)
        if et_als and (target_language == 'tr' or target_language == 'same'):
            audit_findings.append(f"⚠️ **Yabancı Atıf Kalıbı Bulundu ({len(et_als)} adet 'et al.'):** Türkçe metinde İngilizce atıf eki kullanılmış.")
            cleaned_text = re.sub(r'\bet al\.?(?=[\s,;:\)]|$)', 've ark.', cleaned_text)
            corrections_made.append("✅ Tüm 'et al.' yabancı atıf ibareleri Türkçe standartlarına uygun olarak 've ark.' ifadesine dönüştürüldü.")

    # 5. Handle Language Translation Chunk by Chunk
    if target_language in ['tr', 'ar', 'en'] and target_language != 'same':
        target_name = LANG_CONFIGS[target_language]["name"]
        audit_findings.append(f"🌐 **Akademik Çeviri Talebi:** Belgenin tamamı **{target_name}** diline çevriliyor.")
        
        if not llm_func:
            audit_findings.append("⚠️ **Yapay Zeka Servisi Bağlantısı Yok:** Çeviri yapabilmek için sol menüden geçerli bir API Key (Gemini/OpenAI) girmeli veya Ollama çalıştırmalısınız!")
        else:
            # Paragraph-aware chunking: split on blank lines first so a chunk
            # boundary never falls in the middle of a sentence, then group
            # consecutive paragraphs up to ~2000 chars per chunk. Table blocks
            # (consecutive '|' lines) are never split across chunks.
            paragraphs = re.split(r'\n\s*\n', cleaned_text)
            chunks = []
            current_chunk = []
            current_len = 0
            for para in paragraphs:
                para_is_table = para.strip().startswith('|')
                para_len = len(para)
                if current_chunk and current_len + para_len >= 2000 and not para_is_table:
                    chunks.append("\n\n".join(current_chunk))
                    current_chunk = []
                    current_len = 0
                current_chunk.append(para)
                current_len += para_len
            if current_chunk:
                chunks.append("\n\n".join(current_chunk))

            translated_chunks = []
            failed_chunk_count = 0
            system_instr = LANG_CONFIGS[target_language]["system_instruction"]

            def _translate_one_chunk(chunk_text: str, idx: int) -> Optional[str]:
                # Returns translated text, or None if chunk translation fails
                if target_language == "tr":
                    prompt = (
                        f"Aşağıdaki metni (Parça {idx+1}/{len(chunks)}) AKADEMİK TÜRKÇE DİLİNE EKSİKSİZ VE TAM TERCÜME ET. "
                        f"Metin orijinalinde Arapça, İngilizce veya başka bir dilde olsa dahi TÜM PARAGRAFLARI VE CÜMLELERİ TÜRKÇE'YE ÇEVİR.\n\n"
                        f"ÇEVİRİ KURALLARI:\n"
                        f"1. Metnin Tamamını %100 SAF AKADEMİK TÜRKÇE yap. Hiçbir Arapça, İngilizce veya yabancı paragraf/cümle bırakma.\n"
                        f"2. Başlıkların Markdown `#` yapısını koru.\n"
                        f"3. Tabloların Markdown `| Col | Col |` yapısını koru.\n"
                        f"4. YALNIZCA tercüme edilmiş metni döndür. Hiçbir açıklama notu, giriş cümlesi veya yapay zeka gevezeliği ekleme.\n\n"
                        f"TÜRKÇE'YE ÇEVRİLECEK METİN:\n"
                        f"{chunk_text}\n"
                    )
                elif target_language == "ar":
                    prompt = (
                        f"قم بترجمة النص الأكاديمي التالي (الجزء {idx+1}/{len(chunks)}) إلى اللغة العربية الأكاديمية الفصحى الرصينة بنسبة 100%.\n\n"
                        f"قواعد الترجمة الأكاديمية:\n"
                        f"1. ترجم كافة الفقرات والجمل والجداول بالكامل إلى العربية الفصحى الفائقة الجودة.\n"
                        f"2. لا تترك أي جمل بالتركية أو الإنجليزية.\n"
                        f"3. حافظ تماماً على هيكلية العناوين `#` والتنسيق الجدولي `|`.\n"
                        f"4. أعد النص المترجم فقط بدون أي مقدمات أو ملاحظات جانبية.\n\n"
                        f"النص المراد ترجمته:\n"
                        f"{chunk_text}\n"
                    )
                else:
                    prompt = (
                        f"Translate the following text (Part {idx+1}/{len(chunks)}) into 100% ACADEMIC ENGLISH.\n\n"
                        f"TRANSLATION RULES:\n"
                        f"1. Translate all paragraphs and sentences into Academic English.\n"
                        f"2. Preserve Markdown `#` heading hierarchy.\n"
                        f"3. Preserve Markdown `|` tables.\n"
                        f"4. Return ONLY the translated text without any commentary.\n\n"
                        f"TEXT TO TRANSLATE:\n"
                        f"{chunk_text}\n"
                    )
                try:
                    res = llm_func(prompt, system_prompt=system_instr)
                except TypeError:
                    res = llm_func(prompt)
                except Exception as chunk_err:
                    logger.warning(f"Parça {idx+1} çevirisi başarısız: {chunk_err}")
                    return None

                if not res or len(res.strip()) <= 20:
                    return None

                cleaned_res_lines = []
                for l in res.splitlines():
                    if re.match(r'^\s*(Please note|Note:|ملاحظة:|Here is|إليك النص)\b', l, re.IGNORECASE):
                        continue
                    cleaned_res_lines.append(l)
                clean_res = "\n".join(cleaned_res_lines)
                return sanitize_turkish_text(clean_res, language=target_language)

            for idx, chunk in enumerate(chunks):
                if not chunk.strip():
                    continue
                from engine.job_tracker import is_job_cancelled
                if is_job_cancelled():
                    logger.info("Çeviri işlemi kullanıcı tarafından iptal edildi.")
                    break
                if progress_callback:
                    pct = 0.1 + (0.8 * ((idx + 1) / len(chunks)))
                    progress_callback(pct, f"Çeviri Yapılıyor: Parça {idx+1}/{len(chunks)}", f"%{int(pct*100)} tamamlandı")

                # Each chunk gets its own isolated attempt (with one retry) so
                # a single transient API failure doesn't abort the whole
                # translation, and a failed chunk is never silently left
                # untranslated in the final document without a visible marker.
                result = _translate_one_chunk(chunk, idx)
                if result is None:
                    result = _translate_one_chunk(chunk, idx)  # one retry

                if result is not None:
                    translated_chunks.append(result)
                else:
                    failed_chunk_count += 1
                    marker = "⚠️ [BU BÖLÜM OTOMATİK ÇEVİRİLEMEDİ — LÜTFEN ELLE KONTROL EDİN] ⚠️"
                    translated_chunks.append(f"{marker}\n\n{chunk}")

            if translated_chunks:
                cleaned_text = "\n\n".join(translated_chunks)
                if failed_chunk_count:
                    audit_findings.append(
                        f"⚠️ **{failed_chunk_count} parça çevrilemedi:** Bu parçalar orijinal dilinde bırakıldı ve "
                        f"belgede '⚠️ [BU BÖLÜM OTOMATİK ÇEVİRİLEMEDİ...]' etiketiyle işaretlendi — lütfen bulup elle çevirin."
                    )
                    corrections_made.append(f"✅ Belgenin {len(chunks) - failed_chunk_count}/{len(chunks)} parçası profesyonel akademik {target_name} diline çevrildi.")
                else:
                    corrections_made.append(f"✅ Belgenin tamamı ({len(chunks)} parça halinde) profesyonel akademik {target_name} diline çevrildi.")

        # Always run fallback heading translation
        cleaned_text = translate_headings(cleaned_text, target_language)

    default_finding = (
        "✅ لم يتم العثور على أخطاء هيكلية حرجة. تم تطبيق تحسينات التنسيق العامة." if target_language == "ar" else 
        "✅ No critical structural errors found. General formatting improvements applied." if target_language == "en" else 
        "✅ Belgede kritik bir yapısal hata bulunamadı. Genel biçimlendirme iyileştirmeleri uygulandı."
    )
    default_correction = (
        "✅ تم تطبيق جميع تعديلات الفقرات والعناوين بنجاح." if target_language == "ar" else 
        "✅ All paragraph and heading adjustments applied successfully." if target_language == "en" else 
        "✅ Tüm paragraf ve başlık düzenlemeleri başarıyla uygulandı."
    )

    report = {
        "findings": audit_findings or [default_finding],
        "corrections": corrections_made or [default_correction],
        "original_word_count": len(text.split()),
        "cleaned_word_count": len(cleaned_text.split())
    }

    return cleaned_text, report
