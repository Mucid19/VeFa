# -*- coding: utf-8 -*-
"""
VeFa - Kısaltmalar Listesi Modülü
LLM ile tez konusu + bölüm başlıklarından otomatik kısaltma çıkarır ve
YÖK standartlarına uygun Word sayfası üretir (2 kolonlu tablo: Kısaltma | Açılımı).
"""

import re
import logging
from typing import List, Tuple, Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# LLM Prompt — 3 dil
# ---------------------------------------------------------------------------

_ABBR_PROMPT = {
    "tr": """\
Aşağıdaki akademik tez konusu ve tez metninden hareketle, tezde geçen veya geçmesi beklenen tüm \
akademik/teknik kısaltmaları ve açılımlarını Türkçe akademik kurallara göre listele.

TEZ KONUSU: {topic}

TEZ METNİNDEN KISALTMA ADAYLARI:
{text_snippet}

ÇIKTI FORMATI — Her satıra tam olarak şu yapıyı yaz (başka hiçbir şey yazma):
KISALTMA: Açılımı

Örnek:
YÖK: Yükseköğretim Kurulu
APA: American Psychological Association
TDK: Türk Dil Kurumu
DİA: Diyanet İslam Ansiklopedisi
İSAM: İslam Araştırmaları Merkezi
vb.: ve benzeri
bkz.: bakınız
a.g.e.: adı geçen eser
a.g.m.: adı geçen makale
s.: sayfa
C.: cilt
nr.: numara

Yalnızca tezde gerçekten geçen veya konuya özgü kısaltmaları yaz. Alfabetik sırala. Başka açıklama veya not ekleme.

ÖNEMLİ KURALLAR:
1. Satır başlarına madde işareti (-, *, vb.) KESİNLİKLE KOYMA.
2. Kalın (bold) veya eğik (italik) yazı (**, _, vs.) KESİNLİKLE KULLANMA. Doğrudan kısaltmayı yaz.
3. Sadece GERÇEK kısaltmaları (TDK, YÖK, a.g.e., s., vb.) listele. Özel isimleri (Irak, Kûfe, Ahmet, vb.) veya normal kelimeleri (bireysel, yöntem) kısaltma sanıp listeye KESİNLİKLE EKLEME!
""",

    "ar": """\
استخرج جميع الاختصارات الأكاديمية والمصطلحات المختصرة الواردة في الأطروحة التالية أو المتوقع ورودها، \
وأدرج مقابل كل اختصار مدلوله الكامل باللغة العربية.

موضوع الأطروحة: {topic}

مقتطف من نص الأطروحة:
{text_snippet}

صيغة الإخراج — اكتب في كل سطر بالضبط:
الاختصار: المدلول الكامل

مثال:
ص.: صفحة
ج.: جزء
تح.: تحقيق
ط.: طبعة
د.ت.: دون تاريخ
د.م.: دون مكان

رتّب الاختصارات أبجدياً. لا تضف أي تعليق أو ملاحظة خارج الصيغة المطلوبة.\
""",

    "en": """\
Extract all academic abbreviations and acronyms that appear in or are relevant to the following \
thesis, and provide the full expansion for each.

THESIS TOPIC: {topic}

TEXT EXCERPT:
{text_snippet}

OUTPUT FORMAT — write exactly one per line:
ABBR: Full expansion

Example:
APA: American Psychological Association
YÖK: Council of Higher Education (Turkey)
cf.: confer (compare)
ibid.: ibidem (in the same place)
op. cit.: opere citato (in the work cited)
vol.: volume
p.: page
ed.: editor / edition

List only abbreviations actually present or highly relevant to the topic. \
Sort alphabetically. Do not add any other text or commentary.\
"""
}


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

def parse_abbreviation_response(raw: str) -> List[Tuple[str, str]]:
    """
    Parse LLM output lines of the form  'ABBR: Full expansion'
    Returns a deduplicated, alphabetically sorted list of (abbr, expansion) tuples.
    """
    results: dict[str, str] = {}
    for line in raw.splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        m = re.match(r'^[\-\*\s]*(\**)(.+?)\1\s*[:\-–—]\s+(.+)$', line)
        if m:
            abbr = m.group(2).strip(" \t-*_\"'")
            expansion = m.group(3).strip(" \t-*_\"'")
            if 1 <= len(abbr) <= 20 and expansion:
                results[abbr] = expansion

    return sorted(results.items(), key=lambda x: x[0].upper())


# ---------------------------------------------------------------------------
# LLM extractor
# ---------------------------------------------------------------------------

def extract_abbreviations_via_llm(
    topic: str,
    full_markdown: str,
    language: str,
    llm_func,
    system_instruction: str = ""
) -> List[Tuple[str, str]]:
    """
    Call the LLM to extract abbreviations from the thesis text.

    Parameters
    ----------
    topic            : thesis topic string
    full_markdown    : complete generated markdown (used as context)
    language         : "tr" | "ar" | "en"
    llm_func         : callable(prompt, system_prompt=...) → str
    system_instruction: system prompt from the engine's lang_config

    Returns
    -------
    List of (abbr, expansion) tuples, sorted A→Z.
    """
    lang = language if language in _ABBR_PROMPT else "tr"
    prompt_tmpl = _ABBR_PROMPT[lang]

    # Pass up to ~3 000 chars of the markdown as context so the LLM can
    # actually see what abbreviations are in the body.
    snippet = full_markdown[:3000].strip()

    prompt = prompt_tmpl.format(topic=topic, text_snippet=snippet)

    try:
        if system_instruction:
            raw = llm_func(prompt, system_prompt=system_instruction)
        else:
            raw = llm_func(prompt)
    except Exception:
        try:
            raw = llm_func(prompt)
        except Exception as e:
            logger.warning(f"Kısaltma çıkarım LLM hatası: {e}")
            return []

    pairs = parse_abbreviation_response(raw)
    logger.info(f"Kısaltmalar çıkarıldı: {len(pairs)} adet ({language})")
    return pairs


# ---------------------------------------------------------------------------
# Word page builder (called by TurkishDocxGenerator)
# ---------------------------------------------------------------------------

_PAGE_TITLES = {
    "tr": "KISALTMALAR",
    "ar": "قائمة الاختصارات",
    "en": "LIST OF ABBREVIATIONS",
}

_COL_HEADERS = {
    "tr": ("Kısaltma", "Açılımı"),
    "ar": ("الاختصار", "المدلول"),
    "en": ("Abbreviation", "Full Form"),
}


def add_abbreviations_page(
    doc,
    abbr_list: List[Tuple[str, str]],
    language: str = "tr"
) -> None:
    """
    Append a YÖK-compliant Abbreviations page to `doc` (python-docx Document).
    Uses a two-column table: left = abbreviation, right = expansion.
    If abbr_list is empty, inserts a minimal placeholder page.
    """
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    lang = language if language in _PAGE_TITLES else "tr"
    title = _PAGE_TITLES[lang]
    col_abbr, col_exp = _COL_HEADERS[lang]
    is_ar = lang == "ar"

    # --- Page heading ---
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_head.paragraph_format.space_before = Pt(0)
    p_head.paragraph_format.space_after = Pt(18)
    r_head = p_head.add_run(title)
    r_head.font.name = "Times New Roman"
    r_head.font.size = Pt(14)
    r_head.font.bold = True
    r_head.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    if is_ar:
        from engine.turkish_docx_generator import make_paragraph_rtl, make_run_rtl
        make_paragraph_rtl(p_head)
        make_run_rtl(r_head)

    if not abbr_list:
        # Minimal placeholder so the page still exists
        p_ph = doc.add_paragraph()
        p_ph.paragraph_format.space_after = Pt(6)
        r_ph = p_ph.add_run("—")
        r_ph.font.name = "Times New Roman"
        r_ph.font.size = Pt(11)
        doc.add_page_break()
        return

    # --- Two-column table ---
    tbl = doc.add_table(rows=len(abbr_list) + 1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    def _set_cell(cell, text, bold=False, header=False):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        run.font.bold = bold
        if header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            try:
                shd = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="1F4E78" w:color="auto" w:val="clear"/>'
                )
                cell._element.get_or_add_tcPr().append(shd)
            except Exception:
                pass
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Header row
    _set_cell(tbl.rows[0].cells[0], col_abbr, bold=True, header=True)
    _set_cell(tbl.rows[0].cells[1], col_exp, bold=True, header=True)

    # Data rows with zebra stripes
    for idx, (abbr, expansion) in enumerate(abbr_list):
        row = tbl.rows[idx + 1]
        fill = "F2F5FA" if idx % 2 == 0 else "FFFFFF"
        for c_idx, text in enumerate([abbr, expansion]):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10.5)
            if c_idx == 0:
                run.font.bold = True
            try:
                shd = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{fill}" w:color="auto" w:val="clear"/>'
                )
                cell._element.get_or_add_tcPr().append(shd)
            except Exception:
                pass

    # Column widths: abbr ~2 cm, expansion fills the rest
    try:
        for row in tbl.rows:
            row.cells[0].width = Inches(0.8)
            row.cells[1].width = Inches(4.2)
    except Exception:
        pass

    doc.add_page_break()
