# -*- coding: utf-8 -*-
"""
VeFa - Türkçe YÖK Uyumlu Word (.docx) Oluşturucu
Bu modül, Pandoc, WeasyPrint veya LibreOffice gibi harici hiçbir yazılıma ihtiyaç duymadan
doğrudan `python-docx` kütüphanesi kullanarak tam teşekküllü, mükemmel akademik formatta
Türkçe Word (.docx) belgeleri üretir.
"""

import re
import os
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls


def set_cell_background(cell, fill_hex: str):
    """Set shading/background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell internal padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def make_paragraph_rtl(paragraph):
    """Enable Right-to-Left (RTL/bidi) direction on a paragraph for Arabic."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)


def make_run_rtl(run):
    """Enable Right-to-Left (RTL) text rendering on a run for Arabic."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)


def add_page_number_to_footer(run):
    """Insert Word dynamic page number field into a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def set_section_page_numbering(section, fmt="lowerRoman", start=None):
    """Set page number format and optionally restart at a specific number."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))


class TurkishDocxGenerator:
    def __init__(self, metadata: Optional[Dict[str, Any]] = None):
        self.metadata = metadata or {}
        self.doc = Document()
        self.heading_bookmarks = []
        self._bookmark_counter = 0
        self._setup_document_styles()
        from engine.footnotes import FootnoteManager
        self._footnote_manager = FootnoteManager(self.doc)

    def _setup_document_styles(self):
        """Configure page margins and standard academic styles."""
        # Page Size: A4 (21.0 x 29.7 cm) — YÖK zorunlu
        # Margins: Sol 3.0 cm, Sağ/Üst/Alt 2.5 cm — YÖK standardı
        for section in self.doc.sections:
            section.page_width  = Cm(21.0)   # A4 genişlik
            section.page_height = Cm(29.7)   # A4 yükseklik
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3.0)
            section.right_margin  = Cm(2.5)

        # Default Normal style (Times New Roman 12pt, 1.5 line spacing)
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        
        p_format = style.paragraph_format
        p_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p_format.line_spacing = 1.5
        p_format.space_after = Pt(6)
        p_format.space_before = Pt(0)
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_format.first_line_indent = Cm(1.25)  # YÖK: ilk satır girintisi

        # Configure the DOCUMENT'S REAL Heading 1/2/3 styles (not just bold
        # paragraphs styled to look like headings) so Word actually
        # recognizes them: they show up in the Navigation Pane, work with a
        # native Table of Contents field, and appear selected under the
        # correct style name in Word's Styles pane when the cursor is on them.
        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Times New Roman'
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.italic = False
        h1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.keep_with_next = True
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'Times New Roman'
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.font.italic = False
        h2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(6)
        h2.paragraph_format.keep_with_next = True
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        h3 = self.doc.styles['Heading 3']
        h3.font.name = 'Times New Roman'
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.italic = True
        h3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        h3.paragraph_format.space_before = Pt(8)
        h3.paragraph_format.space_after = Pt(4)
        h3.paragraph_format.keep_with_next = True
        h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def create_cover_page(self):
        """Generate YÖK compliant Turkish/Arabic/English Academic Thesis / Paper Cover Page."""
        from engine.turkish_prompts import LANG_CONFIGS
        lang_code = self.metadata.get('language', 'tr').lower()
        cfg = LANG_CONFIGS.get(lang_code, LANG_CONFIGS['tr'])

        author_prefix = cfg.get('author_prefix', 'Hazırlayan')
        advisor_prefix = cfg.get('advisor_prefix', 'Tez Danışmanı')

        institution = self.metadata.get('institution', 'T.C. ÜNİVERSİTESİ').upper()
        faculty = self.metadata.get('faculty', 'LİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ').upper()
        title = self.metadata.get('title', 'AKADEMİK TEZ ÇALIŞMASI').upper()
        author = self.metadata.get('author', 'Yazar Adı Soyadı')
        advisor = self.metadata.get('advisor', 'Danışman Adı Soyadı')
        degree = self.metadata.get('academic_level', 'Yüksek Lisans Tezi')
        year = str(self.metadata.get('year', datetime.now().year))
        city = self.metadata.get('city', 'İstanbul')

        if lang_code == 'ar':
            gov_str = "الجمهورية التركية"
            clean_inst = institution.replace("T.C.", "").replace("TC", "").strip()
            if clean_inst == "ÜNİVERSİTESİ" or not clean_inst:
                clean_inst = "جامعة إسطنبول"
            inst_text = f"{gov_str}\n{clean_inst}"

            if faculty == "LİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ":
                faculty = "معهد الدراسات العليا"

            if degree == "Yüksek Lisans Tezi":
                degree = "أطروحة ماجستير"
            elif degree == "Doktora Tezi":
                degree = "أطروحة دكتوراه"

            if city == "İstanbul":
                city = "إسطنبول"
        else:
            inst_text = f"T.C.\n{institution}"

        # Institution Name
        p1 = self.doc.add_paragraph()
        if lang_code == 'ar':
            make_paragraph_rtl(p1)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(36)
        p1.paragraph_format.space_after = Pt(4)
        run1 = p1.add_run(inst_text)
        if lang_code == 'ar':
            make_run_rtl(run1)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(14)
        run1.font.bold = True

        # Faculty / Institute
        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(72)
        run2 = p2.add_run(faculty)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        run2.font.bold = True

        # Thesis Title
        p3 = self.doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(48)
        run3 = p3.add_run(title)
        run3.font.name = 'Times New Roman'
        run3.font.size = Pt(16)
        run3.font.bold = True
        run3.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)  # Deep Navy Blue

        # Project Type / Degree
        p4 = self.doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_after = Pt(100)
        run4 = p4.add_run(f"({degree.upper()})")
        run4.font.name = 'Times New Roman'
        run4.font.size = Pt(12)
        run4.font.italic = True

        # Author & Advisor Block
        p5 = self.doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5.paragraph_format.space_after = Pt(6)
        r_auth = p5.add_run(f"{author_prefix}: {author}")
        r_auth.font.name = 'Times New Roman'
        r_auth.font.size = Pt(12)
        r_auth.font.bold = True

        if advisor:
            p6 = self.doc.add_paragraph()
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p6.paragraph_format.space_after = Pt(120)
            r_adv = p6.add_run(f"{advisor_prefix}: {advisor}")
            r_adv.font.name = 'Times New Roman'
            r_adv.font.size = Pt(12)

        # Location & Date
        p7 = self.doc.add_paragraph()
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p7.paragraph_format.space_after = Pt(0)
        r_loc = p7.add_run(f"{city} - {year}")
        r_loc.font.name = 'Times New Roman'
        r_loc.font.size = Pt(12)

        # Page Break after Cover Page
        self.doc.add_page_break()

    def create_approval_page(self):
        """
        YÖK 'Onay Sayfası' — jury approval page with signature placeholders.
        Real signatures are added by hand after printing; this only lays out
        the standard template so the student doesn't have to build it from scratch.
        """
        lang_code = self.metadata.get('language', 'tr').lower()
        is_ar = lang_code == 'ar'

        title = self.metadata.get('title', 'AKADEMİK TEZ ÇALIŞMASI').upper()
        author = self.metadata.get('author', 'Yazar Adı Soyadı')
        degree = self.metadata.get('academic_level', 'Yüksek Lisans Tezi')
        institution = self.metadata.get('institution', 'T.C. ÜNİVERSİTESİ')

        if is_ar:
            heading = "صفحة الموافقة"
            body = (
                f"تمت الموافقة على هذه الأطروحة المقدمة من الباحث {author} تحت عنوان "
                f"\"{title}\" باعتبارها مستوفية لمتطلبات درجة {degree} من قبل لجنة المناقشة."
            )
            jury_label = "أعضاء لجنة المناقشة"
            date_label = "التاريخ"
        else:
            heading = "ONAY SAYFASI"
            body = (
                f"{author} tarafından hazırlanan \"{title}\" başlıklı bu tez çalışması "
                f"{institution} tarafından belirlenen jüri üyeleri tarafından incelenerek "
                f"{degree} derecesi için yeterli bulunmuş ve oy birliği/oy çokluğu ile kabul edilmiştir."
            )
            jury_label = "Jüri Üyeleri"
            date_label = "Tez Savunma Tarihi"

        p_head = self.doc.add_paragraph()
        p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_head.paragraph_format.space_after = Pt(24)
        r_head = p_head.add_run(heading)
        r_head.font.name = 'Times New Roman'
        r_head.font.size = Pt(14)
        r_head.font.bold = True
        if is_ar:
            make_paragraph_rtl(p_head)
            make_run_rtl(r_head)

        p_body = self.doc.add_paragraph()
        p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_body.paragraph_format.space_after = Pt(30)
        r_body = p_body.add_run(body)
        r_body.font.name = 'Times New Roman'
        r_body.font.size = Pt(12)
        if is_ar:
            make_paragraph_rtl(p_body)
            make_run_rtl(r_body)

        p_jury = self.doc.add_paragraph()
        p_jury.paragraph_format.space_after = Pt(12)
        r_jury = p_jury.add_run(f"{jury_label}:")
        r_jury.font.name = 'Times New Roman'
        r_jury.font.size = Pt(12)
        r_jury.font.bold = True
        if is_ar:
            make_paragraph_rtl(p_jury)
            make_run_rtl(r_jury)

        for role in (["Danışman", "Üye", "Üye"] if not is_ar else ["المشرف", "عضو", "عضو"]):
            p_row = self.doc.add_paragraph()
            p_row.paragraph_format.space_after = Pt(28)
            r_row = p_row.add_run(f"{role}: _______________________________")
            r_row.font.name = 'Times New Roman'
            r_row.font.size = Pt(12)
            if is_ar:
                make_paragraph_rtl(p_row)
                make_run_rtl(r_row)

        p_date = self.doc.add_paragraph()
        p_date.paragraph_format.space_before = Pt(12)
        r_date = p_date.add_run(f"{date_label}: __ / __ / ____")
        r_date.font.name = 'Times New Roman'
        r_date.font.size = Pt(12)
        if is_ar:
            make_paragraph_rtl(p_date)
            make_run_rtl(r_date)

        self.doc.add_page_break()

    def create_declaration_page(self):
        """
        YÖK 'Bilimsel Etik Beyanı / Tez Beyanı' — a mandatory declaration
        that the thesis is the student's own original work and follows
        academic ethics rules, with a signature line.
        """
        lang_code = self.metadata.get('language', 'tr').lower()
        is_ar = lang_code == 'ar'
        is_en = lang_code == 'en'

        author = self.metadata.get('author', 'Yazar Adı Soyadı')
        title = self.metadata.get('title', 'AKADEMİK TEZ ÇALIŞMASI').upper()
        city = self.metadata.get('city', 'İstanbul')
        year = str(self.metadata.get('year', datetime.now().year))

        if is_ar:
            heading = "إقرار بالأصالة والأخلاقيات العلمية"
            body = (
                f"أقر أنا الموقع أدناه، {author}، بأن أطروحة \"{title}\" التي أعددتها هي من "
                "عملي الأصلي، وأنني التزمت بقواعد الأخلاقيات العلمية والأكاديمية في جميع مراحل "
                "إعدادها، وأنني قمت بالإشارة إلى جميع المصادر التي استفدت منها وفق الأصول "
                "العلمية المتبعة، وأتحمل المسؤولية الكاملة عن أي مخالفة تُكتشف لاحقاً."
            )
            sign_label = "التوقيع"
        elif is_en:
            heading = "DECLARATION OF ACADEMIC HONESTY"
            body = (
                f"I, {author}, hereby declare that the thesis titled \"{title}\" is my own "
                "original work, that I have followed academic and scientific ethics rules "
                "throughout its preparation, and that I have properly cited all sources I "
                "have made use of. I accept full responsibility for any violation discovered "
                "hereafter."
            )
            sign_label = "Signature"
        else:
            heading = "BİLİMSEL ETİK BEYANI"
            body = (
                f"Bu tezin/çalışmanın bana ait, özgün bir çalışma olduğunu; çalışmanın "
                f"hazırlık, veri toplama, analiz ve bilgilerin sunumu aşamalarında bilimsel "
                f"etik ilke ve kurallara uygun davrandığımı; bu çalışma kapsamında bana ait "
                f"olmayan tüm veri, düşünce, görüş ve bilgilere bilimsel etik kuralların "
                f"gereği olarak eksiksiz kaynak gösterdiğimi ve atıfta bulunduğumu, aksinin "
                f"ortaya çıkması durumunda her türlü yasal sonucu kabul ettiğimi beyan ederim."
            )
            sign_label = "İmza"

        p_head = self.doc.add_paragraph()
        p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_head.paragraph_format.space_after = Pt(24)
        r_head = p_head.add_run(heading)
        r_head.font.name = 'Times New Roman'
        r_head.font.size = Pt(14)
        r_head.font.bold = True
        if is_ar:
            make_paragraph_rtl(p_head)
            make_run_rtl(r_head)

        p_body = self.doc.add_paragraph()
        p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_body.paragraph_format.space_after = Pt(48)
        r_body = p_body.add_run(body)
        r_body.font.name = 'Times New Roman'
        r_body.font.size = Pt(12)
        if is_ar:
            make_paragraph_rtl(p_body)
            make_run_rtl(r_body)

        p_sign = self.doc.add_paragraph()
        p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_sign.paragraph_format.space_before = Pt(24)
        r_sign = p_sign.add_run(f"{city}, {year}\n{author}\n{sign_label}: _______________________________")
        r_sign.font.name = 'Times New Roman'
        r_sign.font.size = Pt(12)
        if is_ar:
            p_sign.alignment = WD_ALIGN_PARAGRAPH.LEFT
            make_paragraph_rtl(p_sign)
            make_run_rtl(r_sign)

        self.doc.add_page_break()

    def add_table_of_contents_placeholder(self, headings: List[tuple]):
        """Create a beautifully formatted Academic 'İÇİNDEKİLER / Table of Contents / الفهرس' page."""
        from engine.turkish_prompts import LANG_CONFIGS
        lang_code = self.metadata.get('language', 'tr').lower()
        cfg = LANG_CONFIGS.get(lang_code, LANG_CONFIGS['tr'])

        toc_title = cfg.get('toc_title', 'İÇİNDEKİLER')

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(toc_title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

        # Render headings hierarchy in TOC
        for level, text in headings:
            if level > 3:
                continue
            toc_p = self.doc.add_paragraph()
            toc_p.paragraph_format.space_before = Pt(2)
            toc_p.paragraph_format.space_after = Pt(2)
            toc_p.paragraph_format.line_spacing = 1.2
            toc_p.paragraph_format.first_line_indent = Pt(0)
            
            # Indentation level
            if level == 1:
                toc_p.paragraph_format.left_indent = Inches(0.0)
                run_t = toc_p.add_run(text)
                run_t.font.bold = True
                run_t.font.size = Pt(11)
            elif level == 2:
                toc_p.paragraph_format.left_indent = Inches(0.25)
                run_t = toc_p.add_run(text)
                run_t.font.size = Pt(10.5)
            else:
                toc_p.paragraph_format.left_indent = Inches(0.5)
                run_t = toc_p.add_run(text)
                run_t.font.size = Pt(10)
                run_t.font.italic = True

        self.doc.add_page_break()

    def create_abbreviations_page(self, abbr_list) -> None:
        """
        YÖK Kısaltmalar Listesi sayfası.
        abbr_list: List[Tuple[str, str]]  — (kısaltma, açılım)
        Modülü dışarıdan import ederek çağırır; böylece döngüsel import riski yok.
        """
        from engine.abbreviations import add_abbreviations_page
        lang_code = self.metadata.get('language', 'tr').lower()
        add_abbreviations_page(self.doc, abbr_list, language=lang_code)

    def create_figures_page(self, figures) -> None:
        """
        YÖK Şekiller Listesi sayfası.
        figures: List[str]  — şekil başlıkları
        """
        from engine.figure_table_scanner import add_figures_list_page
        lang_code = self.metadata.get('language', 'tr').lower()
        add_figures_list_page(self.doc, figures, language=lang_code)

    def create_tables_list_page(self, tables) -> None:
        """
        YÖK Tablolar Listesi sayfası.
        tables: List[str]  — tablo başlıkları
        """
        from engine.figure_table_scanner import add_tables_list_page
        lang_code = self.metadata.get('language', 'tr').lower()
        add_tables_list_page(self.doc, tables, language=lang_code)

    def build_from_markdown(
        self,
        markdown_text: str,
        abbr_list=None,
        figures=None,
        tables=None,
    ):
        """Parse complete markdown document into Word paragraphs, headings, and tables.

        abbr_list : List[Tuple[str, str]] | None  — inserted as Kısaltmalar page after TOC
        figures   : List[str] | None              — inserted as Şekiller Listesi page
        tables    : List[str] | None              — inserted as Tablolar Listesi page

        YÖK page order after this method:
            İçindekiler → Kısaltmalar → Şekiller → Tablolar → Metin
        """
        lines = markdown_text.splitlines()
        headings = []
        i = 0

        # Step 1: Collect headings for TOC
        for line in lines:
            line_str = line.strip()
            if line_str.startswith('#'):
                level = len(line_str) - len(line_str.lstrip('#'))
                title = line_str.lstrip('#').strip()
                headings.append((level, title))

        # Add TOC Page
        self.add_table_of_contents_placeholder(headings)

        # Step 1.5: Optional YÖK supplementary list pages
        # (Kısaltmalar → Şekiller → Tablolar)
        if abbr_list is not None:
            self.create_abbreviations_page(abbr_list)
        if figures is not None and len(figures) > 0:
            self.create_figures_page(figures)
        if tables is not None and len(tables) > 0:
            self.create_tables_list_page(tables)

        # Setup Header/Footer for preliminary pages (Roman numerals)
        prelim_section = self.doc.sections[0]
        set_section_page_numbering(prelim_section, fmt="lowerRoman", start=1)
        # YÖK standard: Title page (first page) should not have a page number printed
        prelim_section.different_first_page_header_footer = True

        prelim_footer = prelim_section.footer
        prelim_footer_p = prelim_footer.paragraphs[0] if prelim_footer.paragraphs else prelim_footer.add_paragraph()
        prelim_footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prelim_footer_run = prelim_footer_p.add_run()
        prelim_footer_run.font.name = 'Times New Roman'
        prelim_footer_run.font.size = Pt(10)
        prelim_footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        add_page_number_to_footer(prelim_footer_run)
        
        # Insert Section Break (Next Page) for Main Body
        from docx.enum.section import WD_SECTION
        self.doc.add_section(WD_SECTION.NEW_PAGE)
        
        # Setup Header/Footer for main body (Arabic numerals, restart at 1)
        main_section = self.doc.sections[-1]
        set_section_page_numbering(main_section, fmt="decimal", start=1)
        main_section.different_first_page_header_footer = False
        
        # Break link to previous footer so it doesn't just copy the Roman numeral setup blindly
        main_section.footer.is_linked_to_previous = False
        main_footer = main_section.footer
        main_footer_p = main_footer.paragraphs[0] if main_footer.paragraphs else main_footer.add_paragraph()
        main_footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        main_footer_run = main_footer_p.add_run()
        main_footer_run.font.name = 'Times New Roman'
        main_footer_run.font.size = Pt(10)
        main_footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        add_page_number_to_footer(main_footer_run)

        # Step 2: Parse body content
        while i < len(lines):
            line = lines[i].rstrip()

            # Empty line
            if not line:
                i += 1
                continue

            # Page Break
            if line.strip() in ['---', '***', '___']:
                self.doc.add_page_break()
                i += 1
                continue

            # Check for Table start
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1] and '-' in lines[i + 1]:
                table_lines = []
                # Check for caption above table
                caption = None
                if i > 0 and lines[i - 1].strip().lower().startswith(('tablo', 'table')):
                    caption = lines[i - 1].strip()

                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._render_markdown_table(table_lines, caption)
                continue

            # Headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title_text = line.lstrip('#').strip()
                self._add_heading(title_text, level)
                i += 1
                continue

            # Bullet List (Ensure it's not a heading or keyphrase line)
            if line.strip().startswith(('-', '*', '+ ')) and not line.strip().startswith(('***', '---')):
                clean_item = re.sub(r'^[\-\*\+]\s+', '', line.strip())
                
                # Check if it's actually a Keyphrase line ("Anahtar Kelimeler:" or "Keywords:")
                if clean_item.lower().startswith(('anahtar kelimeler:', 'keywords:')):
                    p = self.doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(12)
                    self._add_formatted_runs(p, clean_item, size_pt=11)
                    i += 1
                    continue

                # Check if it's actually a numbered heading like "* 1.1. Problem Tanımı"
                match_num_heading = re.match(r'^(\d+(?:\.\d+)+)\s+(.*)', clean_item)
                if match_num_heading:
                    h_num = match_num_heading.group(1)
                    h_title = match_num_heading.group(2)
                    level = min(3, h_num.count('.') + 1)
                    self._add_heading(f"{h_num} {h_title}", level)
                    i += 1
                    continue

                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Pt(-12)
                p.paragraph_format.space_after = Pt(3)
                run_bullet = p.add_run("• ")
                run_bullet.font.bold = True
                run_bullet.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
                self._add_formatted_runs(p, clean_item)
                i += 1
                continue

            # Check for section heading patterns (e.g. "1. BÖLÜM: GİRİŞ", "1.1. Problem Tanımı")
            match_sec_heading = re.match(r'^(\d+(?:\.\d+)*\.?)\s+([^\n]+)', line.strip())
            if match_sec_heading and ('bölüm' in line.lower() or len(line.strip()) < 80 and not line.strip().endswith('.')):
                h_num = match_sec_heading.group(1).rstrip('.')
                h_title = match_sec_heading.group(2).strip().rstrip(':')
                dots = h_num.count('.')
                level = 1 if ('bölüm' in line.lower() or dots == 0) else (2 if dots == 1 else 3)
                self._add_heading(line.strip(), level)
                i += 1
                continue

            # Numbered List (Ensure it's a list item, e.g. "1. Metin..." ending in a sentence)
            match_num = re.match(r'^(\d+)[\.\)]\s+(.*)', line.strip())
            if match_num:
                num_str = match_num.group(1)
                item_text = match_num.group(2)
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Pt(-12)
                p.paragraph_format.space_after = Pt(3)
                run_num = p.add_run(f"{num_str}. ")
                run_num.font.bold = True
                run_num.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
                self._add_formatted_runs(p, item_text)
                i += 1
                continue

            # Blockquote
            if line.strip().startswith('>'):
                quote_text = line.strip().lstrip('>').strip()
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.6)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                self._add_formatted_runs(p, quote_text, italic_override=True, size_pt=11)
                i += 1
                continue

            # Regular Paragraph
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.4)  # 1 cm indent for Turkish thesis
            p.paragraph_format.space_after = Pt(6)
            self._add_formatted_runs(p, line.strip())
            i += 1

    def _add_heading(self, text: str, level: int):
        """Add a REAL Word heading (Heading 1/2/3 style) so it shows up in the
        Navigation Pane, works with a native Table of Contents field, and is
        recognized/selected as that style in Word's Styles pane."""
        lang_code = self.metadata.get('language', 'tr').lower()
        style_name = f'Heading {min(max(level, 1), 3)}'
        p = self.doc.add_paragraph(style=style_name)

        if lang_code == 'ar':
            make_paragraph_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run_text = text if lang_code == 'ar' else (text.upper() if level == 1 else text)
        run = p.add_run(run_text)
        if lang_code == 'ar':
            make_run_rtl(run)


    def _add_formatted_runs(self, paragraph, text: str, italic_override=False, size_pt=12):
        """Parse inline Markdown formatting (**bold**, *italic*, `code`) and [[fn:...]] footnote markers into Word runs."""
        lang_code = self.metadata.get('language', 'tr').lower()
        if lang_code == 'ar':
            make_paragraph_rtl(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[\[fn:.*?\]\])', text)
        for token in tokens:
            if not token:
                continue

            # Real Word footnote marker: [[fn:kaynak metni]]
            fn_match = re.match(r'^\[\[fn:(.*?)\]\]$', token, flags=re.DOTALL)
            if fn_match:
                self._footnote_manager.add_footnote(paragraph, fn_match.group(1).strip())
                continue

            run = paragraph.add_run()
            if lang_code == 'ar':
                make_run_rtl(run)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size_pt)

            if italic_override:
                run.font.italic = True

            if token.startswith('**') and token.endswith('**'):
                run.text = token[2:-2]
                run.font.bold = True
            elif token.startswith('*') and token.endswith('*'):
                run.text = token[1:-1]
                run.font.italic = True
            elif token.startswith('`') and token.endswith('`'):
                run.text = token[1:-1]
                run.font.name = 'Courier New'
                run.font.size = Pt(size_pt - 1)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            else:
                run.text = token

    def _render_markdown_table(self, table_lines: List[str], caption: Optional[str] = None):
        """Render markdown pipe table into a native formatted Word table."""
        parsed_rows = []
        for line in table_lines:
            if re.match(r'^\s*\|?\s*:?-+:?\s*\|', line) or '---' in line:  # Skip delimiter line
                continue
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if any(cells):
                parsed_rows.append(cells)

        if not parsed_rows:
            return

        # Table Caption if provided
        if caption:
            cp = self.doc.add_paragraph()
            cp.paragraph_format.space_before = Pt(10)
            cp.paragraph_format.space_after = Pt(4)
            crun = cp.add_run(caption)
            crun.font.name = 'Times New Roman'
            crun.font.size = Pt(10)
            crun.font.bold = True
            crun.font.italic = True

        num_rows = len(parsed_rows)
        num_cols = max(len(r) for r in parsed_rows)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row in enumerate(parsed_rows):
            is_header = (r_idx == 0)
            tr = table.rows[r_idx]

            # Repeat header row on new pages
            if is_header:
                trPr = tr._tr.get_or_add_trPr()
                trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

            for c_idx in range(num_cols):
                cell_text = row[c_idx] if c_idx < len(row) else ""
                cell = tr.cells[c_idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, top=140, bottom=140, left=180, right=180)

                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if not is_header else WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.first_line_indent = Pt(0)

                if is_header:
                    set_cell_background(cell, "1F4E78")  # Deep Navy Blue
                    run = p.add_run(cell_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10.5)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                else:
                    bg_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"
                    set_cell_background(cell, bg_color)
                    self._add_formatted_runs(p, cell_text, size_pt=10)

        # Empty paragraph after table
        sp = self.doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(6)

    def save(self, output_path: str):
        """Save formatted document to destination path."""
        out_p = Path(output_path)
        out_p.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(out_p))
        return str(out_p)


def generate_turkish_academic_docx(
    markdown_text: str,
    output_docx_path: str,
    metadata: Optional[Dict[str, Any]] = None,
    include_cover_page: bool = True,
    include_yok_frontmatter: bool = True,
    abbr_list: Optional[List] = None,
    figures: Optional[List[str]] = None,
    tables: Optional[List[str]] = None,
) -> str:
    """
    Main helper function to convert markdown text to a Turkish Academic Word document.

    include_yok_frontmatter: when True (default) and include_cover_page is True,
        adds the YÖK-required Onay Sayfası (jury approval) and Bilimsel Etik
        Beyanı (declaration of originality) pages right after the cover page,
        before the table of contents.

    abbr_list : List[Tuple[str, str]] | None
        LLM-extracted abbreviation pairs (abbr, expansion). When not None and
        non-empty, a YÖK-compliant Kısaltmalar page is inserted after the TOC.
    figures : List[str] | None
        Figure caption strings scanned from the markdown. When not None and
        non-empty, a Şekiller Listesi page is inserted after abbreviations.
    tables : List[str] | None
        Table caption strings scanned from the markdown. When not None and
        non-empty, a Tablolar Listesi page is inserted after figures.

    YÖK page order:
        Kapak → Onay → Beyan → İçindekiler → Kısaltmalar → Şekiller → Tablolar → Metin
    """
    gen = TurkishDocxGenerator(metadata=metadata)
    if include_cover_page:
        gen.create_cover_page()
        if include_yok_frontmatter:
            gen.create_approval_page()
            gen.create_declaration_page()

    # build_from_markdown inserts the TOC first, then the body.
    # We intercept after TOC and insert the three new list pages.
    gen.build_from_markdown(
        markdown_text,
        abbr_list=abbr_list,
        figures=figures,
        tables=tables,
    )
    return gen.save(output_docx_path)


def enhance_pdf_with_bookmarks(pdf_path: str, headings: List[Tuple[int, str]]) -> bool:
    """
    Inject PDF Outline / Bookmarks Tree (sol navigasyon menüsü) into a PDF file
    using pypdf.
    
    headings: list of (level, title) tuples, e.g. [(1, '1. BÖLÜM: GİRİŞ'), (2, '1.1. Amaç')]
    """
    if not os.path.exists(pdf_path) or not headings:
        return False
    try:
        from pypdf import PdfReader, PdfWriter
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        writer.append(reader)

        parents = {}
        for level, title in headings:
            if level > 3:
                continue
            parent = parents.get(level - 1)
            # Find page index matching title text if possible
            target_page = 0
            for idx_p, page in enumerate(reader.pages):
                txt = page.extract_text() or ""
                if title[:30].lower() in txt.lower():
                    target_page = idx_p
                    break
            
            try:
                bm = writer.add_outline_item(title, page_number=target_page, parent=parent)
                parents[level] = bm
            except Exception:
                pass

        temp_pdf = pdf_path + ".tmp.pdf"
        with open(temp_pdf, "wb") as f_out:
            writer.write(f_out)
        os.replace(temp_pdf, pdf_path)
        logger.info(f"PDF sol menü yer imleri enjekte edildi: {pdf_path}")
        return True
    except Exception as e:
        logger.warning(f"PDF bookmarks enjeksiyonu hatası: {e}")
        return False

