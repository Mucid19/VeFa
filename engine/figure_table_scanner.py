# -*- coding: utf-8 -*-
"""
VeFa - Şekiller ve Tablolar Listesi Tarayıcısı
Üretilen markdown metnini ayrıştırarak tezdeki tablo ve şekil başlıklarını
çeker; YÖK standartlarına uygun Word listesi sayfaları üretir.

Desteklenen başlık kalıpları (Türkçe / Arapça / İngilizce):
  Tablo 1: Karşılaştırma ...
  Table 1: Comparison ...
  جدول 1: مقارنة ...
  Şekil 1: Akış Diyagramı ...
  Figure 1: Flow Chart ...
  شكل 1: مخطط ...
"""

import re
import logging
from typing import List, Tuple

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------

# Matches:  Tablo 1: Başlık  |  Table 1: Title  |  جدول ١:  |  جدول 1:
_TABLE_PATTERN = re.compile(
    r'(?i)^(?:tablo|table|جدول)\s*[\d٠-٩]+[\.:]\s*(.+)',
    re.MULTILINE
)

# Matches:  Şekil 1: Açıklama  |  Figure 1: Caption  |  شكل 1: ...  |  رسم ١:
_FIGURE_PATTERN = re.compile(
    r'(?i)^(?:şekil|sekil|figure|fig\.|شكل|رسم)\s*[\d٠-٩]+[\.:]\s*(.+)',
    re.MULTILINE
)

# Markdown pipe-table header detector (to auto-number unnamed tables)
_PIPE_TABLE_SEP = re.compile(r'^\s*\|?\s*:?-{2,}:?\s*\|', re.MULTILINE)


# ---------------------------------------------------------------------------
# Scanners
# ---------------------------------------------------------------------------

def scan_tables(markdown_text: str) -> List[str]:
    """
    Return a list of table titles found in the markdown.
    Explicit captions (Tablo N: ...) are preferred; if none exist but
    pipe-table separators are found, auto-numbered placeholders are returned.
    """
    explicit = [m.group(1).strip() for m in _TABLE_PATTERN.finditer(markdown_text)]
    if explicit:
        return explicit

    # Fallback: count pipe tables
    sep_count = len(_PIPE_TABLE_SEP.findall(markdown_text))
    if sep_count:
        return [f"Tablo {i + 1}" for i in range(sep_count)]

    return []


def scan_figures(markdown_text: str) -> List[str]:
    """
    Return a list of figure/chart titles found in the markdown.
    Only explicit Şekil/Figure/شكل captions are collected — no fallback
    since figures are optional in text-heavy theses.
    """
    return [m.group(1).strip() for m in _FIGURE_PATTERN.finditer(markdown_text)]


# ---------------------------------------------------------------------------
# Page titles & column headers per language
# ---------------------------------------------------------------------------

_TABLES_TITLE = {
    "tr": "TABLOLAR LİSTESİ",
    "ar": "قائمة الجداول",
    "en": "LIST OF TABLES",
}
_FIGURES_TITLE = {
    "tr": "ŞEKİLLER LİSTESİ",
    "ar": "قائمة الأشكال",
    "en": "LIST OF FIGURES",
}
_TABLE_NUM_HDR = {"tr": "Tablo No", "ar": "الجدول", "en": "Table No."}
_TABLE_CAP_HDR = {"tr": "Başlık", "ar": "العنوان", "en": "Title"}
_FIG_NUM_HDR = {"tr": "Şekil No", "ar": "الشكل", "en": "Figure No."}
_FIG_CAP_HDR = {"tr": "Başlık", "ar": "العنوان", "en": "Title"}


# ---------------------------------------------------------------------------
# Word page builders
# ---------------------------------------------------------------------------

def _add_list_page(
    doc,
    page_title: str,
    items: List[str],
    num_header: str,
    cap_header: str,
    language: str,
) -> None:
    """Generic helper: renders a numbered two-column list page."""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    is_ar = language == "ar"

    # Heading
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_head.paragraph_format.space_before = Pt(0)
    p_head.paragraph_format.space_after = Pt(18)
    r_head = p_head.add_run(page_title)
    r_head.font.name = "Times New Roman"
    r_head.font.size = Pt(14)
    r_head.font.bold = True
    r_head.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    if is_ar:
        from engine.turkish_docx_generator import make_paragraph_rtl, make_run_rtl
        make_paragraph_rtl(p_head)
        make_run_rtl(r_head)

    if not items:
        p_ph = doc.add_paragraph()
        r_ph = p_ph.add_run("—")
        r_ph.font.name = "Times New Roman"
        r_ph.font.size = Pt(11)
        doc.add_page_break()
        return

    tbl = doc.add_table(rows=len(items) + 1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    def _cell(cell, text, bold=False, header=False):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        run.font.bold = bold
        if header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            try:
                shd = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="1F4E78" w:color="auto" w:val="clear"/>'
                )
                cell._element.get_or_add_tcPr().append(shd)
            except Exception:
                pass
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _cell(tbl.rows[0].cells[0], num_header, bold=True, header=True)
    _cell(tbl.rows[0].cells[1], cap_header, bold=True, header=True)

    for idx, title in enumerate(items):
        row = tbl.rows[idx + 1]
        fill = "F2F5FA" if idx % 2 == 0 else "FFFFFF"
        _cell(row.cells[0], str(idx + 1), bold=True)
        _cell(row.cells[1], title)
        for c in range(2):
            try:
                shd = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{fill}" w:color="auto" w:val="clear"/>'
                )
                row.cells[c]._element.get_or_add_tcPr().append(shd)
            except Exception:
                pass

    try:
        for row in tbl.rows:
            row.cells[0].width = Inches(0.8)
            row.cells[1].width = Inches(4.2)
    except Exception:
        pass

    doc.add_page_break()


def add_tables_list_page(doc, tables: List[str], language: str = "tr") -> None:
    """Append a YÖK-compliant Tables List page to `doc`."""
    lang = language if language in _TABLES_TITLE else "tr"
    _add_list_page(
        doc,
        page_title=_TABLES_TITLE[lang],
        items=tables,
        num_header=_TABLE_NUM_HDR[lang],
        cap_header=_TABLE_CAP_HDR[lang],
        language=lang,
    )


def add_figures_list_page(doc, figures: List[str], language: str = "tr") -> None:
    """Append a YÖK-compliant Figures List page to `doc`."""
    lang = language if language in _FIGURES_TITLE else "tr"
    _add_list_page(
        doc,
        page_title=_FIGURES_TITLE[lang],
        items=figures,
        num_header=_FIG_NUM_HDR[lang],
        cap_header=_FIG_CAP_HDR[lang],
        language=lang,
    )
